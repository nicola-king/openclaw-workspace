#!/usr/bin/env python3
"""
Doc Publisher - 功能扩展模块
支持：Word 文档输出、邮件发送、微信发送、网页发布
"""

import os
import sys
from pathlib import Path
from datetime import datetime

class DocPublisherExtensions:
    """文档发布功能扩展"""
    
    def __init__(self):
        self.workspace = Path("/home/nicola/.openclaw/workspace")
        self.output_dir = self.workspace / "published"
        self.output_dir.mkdir(exist_ok=True)
    
    def export_to_word(self, md_file, output_file=None):
        """导出为 Word 文档"""
        print(f"📝 转换 Word: {md_file}")
        
        try:
            import pypandoc
            
            md_path = Path(md_file)
            if output_file is None:
                output_file = str(md_path.with_suffix('.docx'))
            
            output = pypandoc.convert_file(str(md_path), 'docx', outputfile=output_file)
            
            print(f"✅ Word 已生成：{output_file}")
            return True
            
        except ImportError:
            print("⚠️  pypandoc 未安装，使用备用方案...")
            # 备用方案：创建简单的 Word 文档
            return self._create_simple_word(md_file, output_file)
        except Exception as e:
            print(f"❌ 转换失败：{e}")
            return False
    
    def _create_simple_word(self, md_file, output_file):
        """创建简单 Word 文档（备用方案）"""
        try:
            from docx import Document
            
            md_path = Path(md_file)
            content = md_path.read_text(encoding='utf-8')
            
            doc = Document()
            doc.add_heading(md_path.stem, 0)
            
            for line in content.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.strip():
                    doc.add_paragraph(line)
            
            if output_file is None:
                output_file = str(md_path.with_suffix('.docx'))
            
            doc.save(output_file)
            print(f"✅ Word 已生成：{output_file}")
            return True
            
        except Exception as e:
            print(f"❌ Word 生成失败：{e}")
            return False
    
    def send_email(self, pdf_file, subject, recipients, body=None):
        """发送邮件"""
        print(f"📧 发送邮件：{recipients}")
        
        try:
            import smtplib
            from email.mime.multipart import MIMEMultipart
            from email.mime.text import MIMEText
            from email.mime.base import MIMEBase
            from email import encoders
            
            # SMTP 配置（从环境变量读取）
            smtp_server = os.getenv('SMTP_SERVER', 'smtp.gmail.com')
            smtp_port = int(os.getenv('SMTP_PORT', '587'))
            smtp_user = os.getenv('SMTP_USER', '')
            smtp_password = os.getenv('SMTP_PASSWORD', '')
            
            if not smtp_user or not smtp_password:
                print("⚠️  SMTP 配置缺失，请设置环境变量")
                return False
            
            # 创建邮件
            msg = MIMEMultipart()
            msg['From'] = smtp_user
            msg['To'] = ', '.join(recipients)
            msg['Subject'] = subject
            
            # 邮件正文
            if body:
                msg.attach(MIMEText(body, 'plain', 'utf-8'))
            
            # 附加 PDF
            with open(pdf_file, 'rb') as f:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(f.read())
                encoders.encode_base64(part)
                part.add_header(
                    'Content-Disposition',
                    f'attachment; filename={Path(pdf_file).name}'
                )
                msg.attach(part)
            
            # 发送邮件
            server = smtplib.SMTP(smtp_server, smtp_port)
            server.starttls()
            server.login(smtp_user, smtp_password)
            server.send_message(msg)
            server.quit()
            
            print(f"✅ 邮件已发送：{recipients}")
            return True
            
        except Exception as e:
            print(f"❌ 邮件发送失败：{e}")
            return False
    
    def send_wechat(self, pdf_file, title, users=None):
        """发送微信（企业微信）"""
        print(f"💬 发送微信：{users}")
        
        try:
            import requests
            
            # 企业微信配置
            corp_id = os.getenv('WECHAT_CORP_ID', '')
            agent_id = os.getenv('WECHAT_AGENT_ID', '')
            secret = os.getenv('WECHAT_SECRET', '')
            
            if not corp_id or not agent_id or not secret:
                print("⚠️  企业微信配置缺失")
                return False
            
            # 获取 access_token
            token_url = f"https://qyapi.weixin.qq.com/cgi-bin/gettoken?corpid={corp_id}&corpsecret={secret}"
            token_resp = requests.get(token_url)
            token_data = token_resp.json()
            access_token = token_data.get('access_token')
            
            if not access_token:
                print("❌ 获取 access_token 失败")
                return False
            
            # 发送消息
            send_url = f"https://qyapi.weixin.qq.com/cgi-bin/message/send?access_token={access_token}"
            
            # 上传文件
            upload_url = f"https://qyapi.weixin.qq.com/cgi-bin/media/upload?type=file&access_token={access_token}"
            with open(pdf_file, 'rb') as f:
                files = {'media': f}
                upload_resp = requests.post(upload_url, files=files)
                upload_data = upload_resp.json()
                media_id = upload_data.get('media_id')
            
            if not media_id:
                print("❌ 上传文件失败")
                return False
            
            # 发送文件消息
            data = {
                "touser": users or "@all",
                "msgtype": "file",
                "agentid": int(agent_id),
                "file": {"media_id": media_id}
            }
            
            send_resp = requests.post(send_url, json=data)
            send_data = send_resp.json()
            
            if send_data.get('errcode') == 0:
                print(f"✅ 微信已发送：{users}")
                return True
            else:
                print(f"❌ 微信发送失败：{send_data}")
                return False
                
        except Exception as e:
            print(f"❌ 微信发送失败：{e}")
            return False
    
    def publish_to_web(self, md_file, output_dir=None):
        """发布到网页（GitHub Pages）"""
        print(f"🌐 发布网页：{md_file}")
        
        try:
            import markdown2
            
            md_path = Path(md_file)
            if output_dir is None:
                output_dir = self.workspace / "docs"
            
            output_dir = Path(output_dir)
            output_dir.mkdir(exist_ok=True)
            
            # Markdown 转 HTML
            md_content = md_path.read_text(encoding='utf-8')
            html_content = markdown2.markdown(md_content, extras=['tables', 'fenced-code', 'toc'])
            
            # HTML 模板
            html_template = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{md_path.stem}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1 {{ color: #1E88E5; border-bottom: 2px solid #1E88E5; padding-bottom: 10px; }}
        h2 {{ color: #1E88E5; margin-top: 30px; }}
        code {{ background: #f5f5f5; padding: 2px 6px; border-radius: 3px; }}
        pre {{ background: #f5f5f5; padding: 15px; border-radius: 5px; overflow-x: auto; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 12px; text-align: left; }}
        th {{ background: #1E88E5; color: white; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
            
            # 保存 HTML
            html_file = output_dir / f"{md_path.stem}.html"
            html_file.write_text(html_template, encoding='utf-8')
            
            print(f"✅ 网页已发布：{html_file}")
            print(f"🌐 访问地址：file://{html_file.absolute()}")
            return True
            
        except Exception as e:
            print(f"❌ 网页发布失败：{e}")
            return False
    
    def batch_publish(self, input_dir, formats=None):
        """批量发布"""
        print(f"📦 批量发布：{input_dir}")
        
        input_path = Path(input_dir)
        if not input_path.exists():
            print(f"❌ 目录不存在：{input_dir}")
            return False
        
        if formats is None:
            formats = ['pdf', 'docx', 'html']
        
        results = {'pdf': 0, 'docx': 0, 'html': 0}
        
        for md_file in input_path.glob('*.md'):
            print(f"\n处理：{md_file.name}")
            
            if 'pdf' in formats:
                # PDF 已通过 auto-publish-doc.sh 处理
                results['pdf'] += 1
            
            if 'docx' in formats:
                if self.export_to_word(str(md_file)):
                    results['docx'] += 1
            
            if 'html' in formats:
                if self.publish_to_web(str(md_file)):
                    results['html'] += 1
        
        print(f"\n✅ 批量发布完成！")
        print(f"   PDF: {results['pdf']} 个")
        print(f"   Word: {results['docx']} 个")
        print(f"   HTML: {results['html']} 个")
        
        return True


def main():
    """主函数"""
    publisher = DocPublisherExtensions()
    
    if len(sys.argv) < 2:
        print("用法：python3 doc-publisher-extensions.py <命令> [参数]")
        print("\n可用命令:")
        print("  word <文件.md> [输出.docx]  - 导出 Word")
        print("  email <文件.pdf> <主题> <收件人> - 发送邮件")
        print("  wechat <文件.pdf> <标题> [用户] - 发送微信")
        print("  web <文件.md> [输出目录] - 发布网页")
        print("  batch <目录> [格式] - 批量发布")
        sys.exit(1)
    
    command = sys.argv[1]
    
    if command == 'word':
        if len(sys.argv) < 3:
            print("用法：word <文件.md> [输出.docx]")
            sys.exit(1)
        publisher.export_to_word(sys.argv[2], sys.argv[3] if len(sys.argv) > 3 else None)
    
    elif command == 'email':
        if len(sys.argv) < 5:
            print("用法：email <文件.pdf> <主题> <收件人 1,收件人 2,...>")
            sys.exit(1)
        recipients = sys.argv[4].split(',')
        publisher.send_email(sys.argv[2], sys.argv[3], recipients)
    
    elif command == 'wechat':
        if len(sys.argv) < 4:
            print("用法：wechat <文件.pdf> <标题> [用户]")
            sys.exit(1)
        users = sys.argv[4] if len(sys.argv) > 4 else None
        publisher.send_wechat(sys.argv[2], sys.argv[3], users)
    
    elif command == 'web':
        if len(sys.argv) < 3:
            print("用法：web <文件.md> [输出目录]")
            sys.exit(1)
        publisher.publish_to_web(sys.argv[2], sys.argv[3] if len(sys.argv) > 3 else None)
    
    elif command == 'batch':
        if len(sys.argv) < 3:
            print("用法：batch <目录> [格式]")
            sys.exit(1)
        formats = sys.argv[3].split(',') if len(sys.argv) > 3 else None
        publisher.batch_publish(sys.argv[2], formats)
    
    else:
        print(f"❌ 未知命令：{command}")
        sys.exit(1)


if __name__ == "__main__":
    main()
